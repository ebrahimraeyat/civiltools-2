"""
civiltools.wind.report
======================
Word (DOCX) report generator for billboard wind load calculations
per Iranian National Building Code – Section 6 (مبحث ششم), Chapter 10.

Requires: python-docx
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from civiltools.wind.billboard import BillboardInputs, WindLoadOutput, WIND_SPEEDS


# ── Colour palette ────────────────────────────────────────────────────────────
_HEADER_BG  = RGBColor(30,  60, 120)   # dark blue — section headers
_ROW_BG_ALT = RGBColor(235, 241, 250)  # very light blue — alternating table rows
_ACCENT     = RGBColor(50,  100, 50)   # green — numeric results
_BLACK      = RGBColor(0,   0,  0)

# Hex equivalents for cell shading (RGBColor has no .red/.green/.blue in python-docx)
_HEADER_BG_HEX  = "1E3C78"
_ROW_BG_ALT_HEX = "EBF1FA"


# ══════════════════════════════════════════════════════════════════════════════
# Low-level helpers
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell background with a given hex colour string (e.g. '1E3C78')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_paragraph(
    doc: Document,
    text: str,
    style: str = "Normal",
    bold: bool = False,
    italic: bool = False,
    size: int = 11,
    color: Optional[RGBColor] = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    space_after: int = 4,
) -> None:
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, title: str) -> None:
    """Add a visually distinct section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {title}  ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Background colour via paragraph shading (approximate via table trick)
    # Insert a 1×1 table instead for reliable background shading
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _HEADER_BG_HEX)
    cell.width = Inches(6.0)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"  {title}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(255, 255, 255)


def _add_kv_table(
    doc: Document,
    rows: list[tuple[str, str]],
) -> None:
    """
    Add a two-column key-value table with alternating row shading.

    Parameters
    ----------
    doc : Document
    rows : list of (label, value) string pairs
    """
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(rows):
        row = tbl.rows[i]
        row.height = Cm(0.7)

        # Label cell
        lc = row.cells[0]
        lc.width = Cm(8)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lrun = lp.add_run(label)
        lrun.font.size = Pt(10)
        lrun.bold = True

        # Value cell
        vc = row.cells[1]
        vc.width = Cm(8)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vc.paragraphs[0]
        vrun = vp.add_run(value)
        vrun.font.size = Pt(10)
        vrun.font.color.rgb = _ACCENT

        if i % 2 == 1:
            _set_cell_bg(lc, _ROW_BG_ALT_HEX)
            _set_cell_bg(vc, _ROW_BG_ALT_HEX)

    doc.add_paragraph()  # spacing after table


def _add_equation_table(
    doc: Document,
    steps: list[tuple[str, str]],
) -> None:
    """
    Add a calculation-steps table with step descriptions and formulas.

    Parameters
    ----------
    steps : list of (step_label, equation_string)
    """
    tbl = doc.add_table(rows=1 + len(steps), cols=2)
    tbl.style = "Table Grid"

    # Header row
    hrow = tbl.rows[0]
    for j, hdr in enumerate(["Calculation Step", "Value / Formula"]):
        hc = hrow.cells[j]
        _set_cell_bg(hc, _HEADER_BG_HEX)
        hp = hc.paragraphs[0]
        hr_run = hp.add_run(hdr)
        hr_run.bold = True
        hr_run.font.size = Pt(10)
        hr_run.font.color.rgb = RGBColor(255, 255, 255)

    for i, (label, formula) in enumerate(steps):
        row = tbl.rows[i + 1]
        row.height = Cm(0.75)

        lc = row.cells[0]
        lc.width = Cm(8)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lrun = lc.paragraphs[0].add_run(label)
        lrun.font.size = Pt(10)

        vc = row.cells[1]
        vc.width = Cm(8)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vrun = vc.paragraphs[0].add_run(formula)
        vrun.font.size = Pt(10)
        vrun.bold = True
        vrun.font.color.rgb = _ACCENT

        if i % 2 == 1:
            _set_cell_bg(lc, _ROW_BG_ALT_HEX)
            _set_cell_bg(vc, _ROW_BG_ALT_HEX)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# Main report function
# ══════════════════════════════════════════════════════════════════════════════

def generate_word_report(
    inputs: BillboardInputs,
    output: WindLoadOutput,
    save_path: Optional[Path | str] = None,
    sketch_image_path: Optional[str] = None,
) -> Document:
    """
    Generate a styled Word (.docx) report for a billboard wind load calculation.

    Parameters
    ----------
    inputs : BillboardInputs
        The user inputs used for calculation.
    output : WindLoadOutput
        The calculation results as returned by ``calculate_wind_load``.
    save_path : Path or str, optional
        If provided, the document is saved to this path.  The parent
        directory is created if it does not exist.

    Returns
    -------
    docx.Document
        The generated Word document object (can be saved or further modified).
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    tr = title_p.add_run("Wind Load Calculation Report — Free-Standing Billboard")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = _HEADER_BG

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(2)
    sr = subtitle_p.add_run(
        "Iranian National Building Code – Section 6 (مبحث ششم), Chapter 10 (Wind Load)"
    )
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(80, 80, 80)

    import datetime
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(12)
    dr = date_p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── Schematic sketch ───────────────────────────────────────────────────
    if sketch_image_path and Path(sketch_image_path).exists():
        _add_section_heading(doc, "Billboard Schematic")
        doc.add_paragraph()
        sketch_p = doc.add_paragraph()
        sketch_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = sketch_p.add_run()
        run_img.add_picture(sketch_image_path, width=Cm(10))
        cap_p = doc.add_paragraph(
            f"Figure: Side-view schematic — h={inputs.height:.2f} m, "
            f"w={inputs.width:.2f} m, b={inputs.bottom_elevation:.2f} m"
        )
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap_p.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(90, 90, 90)
        doc.add_paragraph()

    # ── Section 1: Input parameters ───────────────────────────────────────────
    _add_section_heading(doc, "1. Input Parameters")
    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Billboard Height (h)", f"{inputs.height:.3f} m"),
        ("Billboard Width (w)", f"{inputs.width:.3f} m"),
        ("Billboard Area (A = h × w)", f"{output.A:.3f} m²"),
        ("Bottom Elevation", f"{inputs.bottom_elevation:.3f} m"),
        ("City", inputs.city),
        ("Terrain Type", inputs.terrain_type.title()),
        ("Support Type", inputs.support_type.replace("_", " ").title()
         + ("  (b = 0 → auto)"
            if inputs.bottom_elevation == 0
            else f"  (b = {inputs.bottom_elevation:.2f} m → auto)")),
        ("Support Length l (= width w)", f"{inputs.width:.3f} m"),
        ("l / h ratio", f"{output.lh_ratio:.4f}"),
        ("Importance Factor (I_w)", f"{inputs.importance_factor:.2f}"),
        ("Topographic Factor (C_t)", f"{inputs.topographic_factor:.2f}"),
    ])

    # ── Section 2: Step-by-step calculations ─────────────────────────────────
    _add_section_heading(doc, "2. Step-by-Step Calculations")
    doc.add_paragraph()

    _add_equation_table(doc, [
        (
            "Step 1a — Basic wind speed (km/h)\n"
            f"  City: {inputs.city}  [Table 1-10-6, pp. 116-117]",
            f"V = {output.V_kmh:.2f} km/h",
        ),
        (
            "Step 1b — Convert to m/s\n"
            "  V_ms = V_kmh / 3.6",
            f"V = {output.V_ms:.4f} m/s",
        ),
        (
            "Step 1c — Basic wind pressure  [Clause 3-10-6]\n"
            "  q = 0.001637 × V_ms²",
            f"q = 0.001637 × ({output.V_ms:.4f})²\n  = {output.q:.4f} kN/m²",
        ),
        (
            "Step 2 — Reference height  [centroid of sign]\n"
            f"  Z = bottom_elevation + h/2\n"
            f"  Z = {inputs.bottom_elevation:.2f} + {inputs.height:.2f}/2",
            f"Z = {output.Z_ref:.4f} m",
        ),
        (
            f"Step 3 — Exposure coefficient Ce  [Clause 6-10-6]\n"
            f"  Terrain: {inputs.terrain_type}\n"
            + (
                f"  Ce = (Z/10)^0.28 = ({output.Z_ref:.2f}/10)^0.28"
                if inputs.terrain_type == "open"
                else f"  Ce = 0.5×(Z/12.7)^0.28 = 0.5×({output.Z_ref:.2f}/12.7)^0.28"
            ),
            f"Ce = {output.Ce:.4f}",
        ),
        (
            "Step 4 — Gust factor Cg  [Clause 8/9-10-6]\n"
            f"  {output.cg_method}",
            f"Cg = {output.Cg:.2f}",
        ),
        (
            f"Step 5 — Force coefficient Cf  [Figure P-6-4-5, p. 156]\n"
            f"  Support: {inputs.support_type},  l/h = {output.lh_ratio:.4f}",
            f"Cf = {output.Cf:.4f}",
        ),
        ("Step 6 — Total wind force  [p. 156]\n"
            "  F = I_w × Cf × q × Cg × Ce × A\n"
            f"  F = {inputs.importance_factor:.2f} × {output.Cf:.4f} × "
            f"{output.q:.4f} × {output.Cg:.2f} × {output.Ce:.4f} × {output.A:.4f}",
            f"F = {output.F_total_kN:.4f} kN\n"
            f"  = {output.F_total_kN * 101.972:.2f} kgf",
        ),
        (
            "Step 7 — Equivalent design pressure\n"
            "  P_design = F / A",
            f"P = {output.P_design_kPa:.4f} kN/m\u00b2\n"
            f"  = {output.P_design_kPa * 101.972:.2f} kgf/m\u00b2",
        ),
    ])

    # ── Section 3: Results summary ────────────────────────────────────────────
    _add_section_heading(doc, "3. Results Summary")
    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Basic Wind Speed", f"{output.V_ms:.2f} m/s  ({output.V_kmh:.2f} km/h)"),
        ("Basic Wind Pressure (q)", f"{output.q:.4f} kN/m²"),
        ("Reference Height (Z)", f"{output.Z_ref:.4f} m"),
        ("Exposure Coefficient (Ce)", f"{output.Ce:.4f}"),
        ("Gust Effect Factor (Cg)", f"{output.Cg:.2f}"),
        ("Force Coefficient (Cf)", f"{output.Cf:.4f}"),
        ("Billboard Area (A)", f"{output.A:.3f} m²"),
        ("Total Wind Force (F)",
         f"{output.F_total_kN:.4f} kN  =  {output.F_total_kN * 101.972:.2f} kgf"),
        ("Design Pressure (P_design)",
         f"{output.P_design_kPa:.4f} kN/m\u00b2  =  {output.P_design_kPa * 101.972:.2f} kgf/m\u00b2"),
    ])

    # ── Section 4: Code references ────────────────────────────────────────────
    _add_section_heading(doc, "4. Code References")
    doc.add_paragraph()
    refs = [
        "Iranian National Building Code, Section 6 (مبحث ششم): Loads on Buildings.",
        "Chapter 10 – Wind Load, Clauses 3-10-6, 6-10-6, 8-10-6, 9-10-6.",
        "Appendix 6-4 – Dynamic analysis methods (pages 146-148).",
        "Figure P-6-4-5 and tables on pages 155-156: Force coefficients for "
        "walls, self-standing panels, and billboards.",
        "Table 1-10-6: Basic wind speed by city (pages 116-117).",
        "Flowchart: Step-by-step wind load calculation procedure (page 169).",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_p.add_run(
        "Generated by civilTools  |  Wind Load Module  |  مبحث ششم – فصل ۱۰"
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(130, 130, 130)

    # ── Save ──────────────────────────────────────────────────────────────────
    if save_path is not None:
        save_path = Path(save_path)
        save_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(save_path))

    return doc

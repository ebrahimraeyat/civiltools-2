"""
Result display widget — ported from civilTools/table_model.py ResultWidget.

Shows a pandas-backed table with:
- Column filter (combobox + text search)
- Sort by clicking headers
- Export to Excel, Word, CSV
- Copy to clipboard
- Color-coded cells via the model's BackgroundRole
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Callable

import pandas as pd
from PySide6 import QtCore
from PySide6.QtCore import QEvent, Qt, QRegularExpression, Signal, QTimer
from PySide6.QtGui import QColor, QKeySequence, QShortcut
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QComboBox, QPushButton, QCheckBox, QTableView, QHeaderView,
    QFileDialog, QMessageBox, QApplication, QSplitter, QTextEdit, QFrame,
    QAbstractItemView, QListWidget, QListWidgetItem, QMenu, QWidgetAction,
)

from civiltools.gui.table_models import PandasModel


class HeaderFilterListWidget(QListWidget):
    """List widget with drag select/deselect and single-click activation."""

    single_item_activated = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSelectionMode(QAbstractItemView.SelectionMode.MultiSelection)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self._press_row: int | None = None
        self._last_row: int | None = None
        self._drag_select_state = True
        self._did_drag = False
        self._press_modifiers = Qt.KeyboardModifier.NoModifier
        self._pending_single_collapse = False

    def _row_at(self, position) -> int:
        item = self.itemAt(position)
        if item is None:
            return -1
        return self.row(item)

    def _apply_range_state(self, start_row: int, end_row: int, selected: bool) -> None:
        for row in range(min(start_row, end_row), max(start_row, end_row) + 1):
            item = self.item(row)
            if item is not None:
                item.setSelected(selected)

    def mousePressEvent(self, event):
        if event.button() != Qt.MouseButton.LeftButton:
            super().mousePressEvent(event)
            return

        self._press_modifiers = event.modifiers()
        if self._press_modifiers & (
            Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier
        ):
            super().mousePressEvent(event)
            return

        press_row = self._row_at(event.position().toPoint())
        if press_row < 0:
            self.clearSelection()
            self._press_row = None
            self._last_row = None
            self._did_drag = False
            super().mousePressEvent(event)
            return

        self._press_row = press_row
        self._last_row = press_row
        self._did_drag = False
        self._pending_single_collapse = False

        item = self.item(press_row)
        if item.isSelected() and len(self.selectedIndexes()) > 1:
            self._pending_single_collapse = True
            self._drag_select_state = False
        else:
            self._drag_select_state = True
            self.clearSelection()
            item.setSelected(True)

        self.setCurrentRow(press_row)
        event.accept()

    def mouseMoveEvent(self, event):
        if (
            self._press_row is None
            or not (event.buttons() & Qt.MouseButton.LeftButton)
            or self._press_modifiers
            & (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier)
        ):
            super().mouseMoveEvent(event)
            return

        move_row = self._row_at(event.position().toPoint())
        if move_row < 0:
            super().mouseMoveEvent(event)
            return

        if self._last_row is None:
            self._last_row = self._press_row
        if move_row != self._last_row:
            if self._pending_single_collapse:
                self._pending_single_collapse = False
            self._apply_range_state(self._last_row, move_row, self._drag_select_state)
            self._did_drag = True
            self._last_row = move_row
        event.accept()

    def mouseReleaseEvent(self, event):
        if event.button() != Qt.MouseButton.LeftButton:
            super().mouseReleaseEvent(event)
            return

        if self._press_row is None:
            super().mouseReleaseEvent(event)
            return

        selected_value = None
        if self._pending_single_collapse:
            self.clearSelection()
            item = self.item(self._press_row)
            if item is not None:
                item.setSelected(True)
                self.setCurrentRow(self._press_row)

        should_auto_apply = (
            not self._did_drag
            and not (
                self._press_modifiers
                & (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.ShiftModifier)
            )
            and len(self.selectedItems()) == 1
        )
        if should_auto_apply:
            selected_value = self.selectedItems()[0].data(Qt.ItemDataRole.UserRole)

        self._press_row = None
        self._last_row = None
        self._did_drag = False
        self._press_modifiers = Qt.KeyboardModifier.NoModifier
        self._pending_single_collapse = False
        event.accept()

        if selected_value is not None:
            self.single_item_activated.emit(str(selected_value))


class ResultWidget(QWidget):
    """Main result display — mirrors civilTools ``table_model.ResultWidget``."""

    #: Emitted when a table row is selected; carries the source DataFrame row index.
    selection_changed = Signal(int)

    def __init__(
        self,
        df: pd.DataFrame,
        model_class: type[PandasModel] = PandasModel,
        delegate_class: type | None = None,
        legend_items: list[tuple[str, str]] | None = None,
        sortable: bool = True,
        cell_selected: Callable[[int, int], None] | None = None,
        title: str = "",
        summary: str = "",
        ok: bool = True,
        parent: QWidget | None = None,
        kwargs: dict | None = None,
    ):
        super().__init__(parent)
        self._title = title
        self._df = df
        self._cell_selected = cell_selected
        # Defer ETABS selection so double-click can open the editor without
        # waiting on a blocking COM call from the first half of the double-click.
        self._pending_cell: tuple[int, int] | None = None
        self._select_timer = QTimer(self)
        self._select_timer.setSingleShot(True)
        # Slightly past the OS double-click threshold so a real double-click
        # never pays for ETABS selection before the editor opens.
        try:
            interval = int(QApplication.doubleClickInterval()) + 50
        except Exception:
            interval = 350
        self._select_timer.setInterval(max(interval, 300))
        self._select_timer.timeout.connect(self._emit_pending_cell_selection)

        # ── Layout ──────────────────────────────────────────────────
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(4, 4, 4, 4)

        # Title
        if title:
            title_lbl = QLabel(f"<b style='font-size:14px'>{title}</b>")
            main_layout.addWidget(title_lbl)

        # Filter bar (matches original: Filter [text] By Column: [combo])
        filter_bar = QHBoxLayout()
        filter_bar.addWidget(QLabel("Filter"))
        self._filter_edit = QLineEdit()
        self._filter_edit.setPlaceholderText("Type to filter…")
        filter_bar.addWidget(self._filter_edit)
        filter_bar.addWidget(QLabel("By Column:"))
        self._col_combo = QComboBox()
        filter_bar.addWidget(self._col_combo)

        # Export buttons
        self._btn_excel = QPushButton("Excel")
        self._btn_excel.setToolTip("Export to Excel (.xlsx)")
        filter_bar.addWidget(self._btn_excel)
        self._btn_word = QPushButton("Word")
        self._btn_word.setToolTip("Export to Word (.docx)")
        filter_bar.addWidget(self._btn_word)
        self._btn_csv = QPushButton("CSV")
        self._btn_csv.setToolTip("Export to CSV")
        filter_bar.addWidget(self._btn_csv)
        self._btn_copy = QPushButton("Copy")
        self._btn_copy.setToolTip("Copy table to clipboard")
        filter_bar.addWidget(self._btn_copy)

        self._open_after = QCheckBox("Open")
        self._open_after.setChecked(True)
        filter_bar.addWidget(self._open_after)

        main_layout.addLayout(filter_bar)

        # Table + detail pane in a vertical splitter
        self._splitter = QSplitter(Qt.Orientation.Vertical)

        self._table = QTableView()
        self._table.setSortingEnabled(sortable)
        self._table.setAlternatingRowColors(True)
        self._table.setSelectionBehavior(QTableView.SelectionBehavior.SelectItems)
        self._table.setSelectionMode(QTableView.SelectionMode.SingleSelection)
        self._splitter.addWidget(self._table)

        # Bottom pane: a horizontal QSplitter that holds
        #   [detail text | (optional right panel added via add_right_panel())]
        self._bottom_splitter = QSplitter(Qt.Orientation.Horizontal)
        self._bottom_splitter.hide()  # shown only when __detail__ column is present

        self._detail_pane = QTextEdit()
        self._detail_pane.setReadOnly(True)
        self._detail_pane.setPlaceholderText("Select a row to see calculation details…")
        self._detail_pane.setFontFamily("Courier New")
        self._bottom_splitter.addWidget(self._detail_pane)

        self._splitter.addWidget(self._bottom_splitter)
        self._splitter.setStretchFactor(0, 3)
        self._splitter.setStretchFactor(1, 1)

        content_layout = QHBoxLayout()
        content_layout.addWidget(self._splitter, 1)
        if legend_items:
            content_layout.addWidget(self._build_color_legend(legend_items))
        main_layout.addLayout(content_layout, 1)

        # Summary bar
        self._summary_lbl = QLabel()
        main_layout.addWidget(self._summary_lbl)

        # ── Model + Proxy ───────────────────────────────────────────
        self._model = model_class(df, kwargs)
        self._proxy = ColumnValueFilterProxyModel(self)
        self._proxy.setSourceModel(self._model)
        self._proxy.setFilterCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self._table.setModel(self._proxy)
        self._table.viewport().installEventFilter(self)
        if delegate_class is not None:
            # Match FreeCAD ControlColumnResultWidget: custom delegate, no sorting.
            self._table.setItemDelegate(delegate_class(self._table))
            self._table.setEditTriggers(
                QAbstractItemView.EditTrigger.DoubleClicked
                | QAbstractItemView.EditTrigger.EditKeyPressed
            )
        # FreeCAD called resizeColumnsToContents() once — do not keep
        # ResizeToContents mode (recalculates on every paint and freezes large tables).
        header = self._table.horizontalHeader()
        header.setSectionResizeMode(
            QHeaderView.ResizeMode.Interactive
        )
        header.setStretchLastSection(False)
        header.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        header.customContextMenuRequested.connect(self._show_header_filter_menu)
        self._table.resizeColumnsToContents()
        if self._model.columnCount():
            last_column = self._model.columnCount() - 1
            self._table.setColumnWidth(
                last_column, min(self._table.columnWidth(last_column), 300)
            )

        # Populate column combo (skip hidden system columns)
        for source_column, col_name in enumerate(self._model.df.columns):
            if not str(col_name).startswith("__"):
                self._col_combo.addItem(str(col_name), source_column)

        # Hide any __ system columns from the visible table
        for source_column, col_name in enumerate(self._model.df.columns):
            if not str(col_name).startswith("__"):
                continue
            self._table.setColumnHidden(source_column, True)

        # If a __detail__ column exists: show detail pane
        self._detail_col = "__detail__"
        if self._detail_col in self._model.df.columns:
            self._table.setSelectionMode(QTableView.SelectionMode.SingleSelection)
            self._bottom_splitter.show()
            # connect after model is set (done below)

        # Summary
        if summary:
            color = "#006400" if ok else "#8B0000"
            self._summary_lbl.setText(
                f'<span style="color:{color}; font-weight:bold">{summary}</span>'
            )

        # ── Connections ─────────────────────────────────────────────
        self._filter_edit.textChanged.connect(self._on_filter_changed)
        self._col_combo.currentIndexChanged.connect(self._on_col_changed)
        self._btn_excel.clicked.connect(self.export_to_excel)
        self._btn_word.clicked.connect(self.export_to_word)
        self._btn_csv.clicked.connect(self.export_to_csv)
        self._btn_copy.clicked.connect(self.copy_to_clipboard)

        # Ctrl+C shortcut
        sc = QShortcut(QKeySequence.StandardKey.Copy, self._table)
        sc.activated.connect(self.copy_to_clipboard)

        # Connect row-click detail pane (after model is assigned)
        if self._detail_col in self._model.df.columns:
            self._table.selectionModel().selectionChanged.connect(self._on_row_selected)
        # FreeCAD wired clicked → show_frame. We keep that behaviour for a true
        # single-click, but cancel it when the user double-clicks to edit.
        if self._cell_selected is not None:
            self._table.clicked.connect(self._on_cell_clicked)
            self._table.doubleClicked.connect(self._on_cell_double_clicked)

    def _build_color_legend(self, legend_items: list[tuple[str, str]]) -> QWidget:
        """Build a compact sidebar explaining table background colors."""
        panel = QFrame()
        panel.setFrameShape(QFrame.Shape.StyledPanel)
        panel.setFixedWidth(190)
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        title = QLabel("Color legend")
        title.setStyleSheet("font-weight: bold;")
        layout.addWidget(title)
        for color, label in legend_items:
            row = QHBoxLayout()
            swatch = QLabel()
            swatch.setFixedSize(16, 16)
            swatch.setStyleSheet(f"background-color: {color}; border: 1px solid #666;")
            row.addWidget(swatch)
            row.addWidget(QLabel(label), 1)
            layout.addLayout(row)
        layout.addStretch()
        return panel

    # ── Public API for extending the bottom panel ───────────────────

    def add_right_panel(self, widget: QWidget) -> None:
        """
        Add a widget to the right of the detail text pane.
        The bottom splitter is shown automatically.
        """
        self._bottom_splitter.addWidget(widget)
        self._bottom_splitter.setStretchFactor(0, 1)  # detail text
        self._bottom_splitter.setStretchFactor(
            self._bottom_splitter.count() - 1, 1
        )  # right panel
        self._bottom_splitter.show()

    # ── Detail pane slot ────────────────────────────────────────────

    def _on_row_selected(self):
        indexes = self._table.selectionModel().selectedRows()
        if not indexes:
            self._detail_pane.clear()
            return
        source_idx = self._proxy.mapToSource(indexes[0])
        df_row = source_idx.row()
        detail_text = self._model.df.iloc[df_row].get(self._detail_col, "")
        self._detail_pane.setPlainText(str(detail_text))
        self.selection_changed.emit(df_row)

    def _on_cell_clicked(self, proxy_index):
        """Queue ETABS selection; cancelled if a double-click follows."""
        if self._cell_selected is None:
            return
        # Never call COM while an in-cell editor is open.
        if self._table.state() is not QAbstractItemView.State.NoState:
            self._select_timer.stop()
            self._pending_cell = None
            return
        source_index = self._proxy.mapToSource(proxy_index)
        self._pending_cell = (source_index.row(), source_index.column())
        self._select_timer.start()

    def _on_cell_double_clicked(self, _proxy_index):
        """Open editor without paying the ETABS selection cost first."""
        self._select_timer.stop()
        self._pending_cell = None

    def _emit_pending_cell_selection(self):
        if self._cell_selected is None or self._pending_cell is None:
            return
        if self._table.state() is not QAbstractItemView.State.NoState:
            self._pending_cell = None
            return
        row, col = self._pending_cell
        self._pending_cell = None
        self._cell_selected(row, col)

    def eventFilter(self, watched, event):
        if (
            watched is self._table.viewport()
            and event.type() == QEvent.Type.MouseButtonRelease
            and event.button() == Qt.MouseButton.LeftButton
        ):
            self._refresh_from_etabs(0, 0)
        return super().eventFilter(watched, event)

    def _refresh_from_etabs(self, row: int, col: int) -> None:
        refresh_from_etabs = getattr(self._model, "refresh_from_etabs", None)
        if refresh_from_etabs is not None:
            refresh_from_etabs(row, col)

    # ── Filter slots ────────────────────────────────────────────────

    def _on_filter_changed(self, text: str):
        regex = QRegularExpression(
            text, QRegularExpression.PatternOption.CaseInsensitiveOption
        )
        self._proxy.setFilterRegularExpression(regex)

    def _on_col_changed(self, index: int):
        source_column = self._col_combo.itemData(index)
        if source_column is not None:
            self._proxy.setFilterKeyColumn(source_column)

    @staticmethod
    def _collect_unique_header_values(source_model, column: int) -> list[str]:
        values = set()
        for row in range(source_model.rowCount()):
            index = source_model.index(row, column)
            value = source_model.data(index, Qt.ItemDataRole.DisplayRole)
            values.add("" if value is None else str(value).strip())
        return sorted(values, key=lambda item: (item == "", item.lower()))

    def _show_header_filter_menu(self, pos) -> None:
        """Show a multi-value filter for the column under the table header."""
        header = self._table.horizontalHeader()
        column = header.logicalIndexAt(pos)
        if column < 0 or self._table.isColumnHidden(column):
            return

        source_model = self._proxy.sourceModel()
        if source_model is None:
            return

        column_title = source_model.headerData(
            column, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole
        )
        if column_title is None or str(column_title).startswith("__"):
            return

        unique_values = self._collect_unique_header_values(source_model, column)
        active_values = self._proxy.get_column_filter_values(column)

        menu = QMenu(self._table)
        title_action = menu.addAction(f"Filter: {column_title}")
        title_action.setEnabled(False)

        if not unique_values:
            empty_action = menu.addAction("(No values)")
            empty_action.setEnabled(False)
        else:
            filter_widget = QWidget(menu)
            filter_layout = QVBoxLayout(filter_widget)
            filter_layout.setContentsMargins(8, 6, 8, 6)
            filter_layout.setSpacing(6)

            value_list = HeaderFilterListWidget(filter_widget)
            for value in unique_values:
                item = QListWidgetItem("{blanks}" if value == "" else value)
                item.setData(Qt.ItemDataRole.UserRole, value)
                value_list.addItem(item)
                if value in active_values:
                    item.setSelected(True)

            row_height = value_list.sizeHintForRow(0) if value_list.count() else 20
            visible_rows = min(max(4, value_list.count()), 8)
            value_list.setFixedHeight((row_height * visible_rows) + 8)

            apply_button = QPushButton("Apply Multiple Item Filter", filter_widget)

            def apply_single_value(selected_value: str) -> None:
                self._proxy.set_column_filter_values(column, {selected_value})

            def apply_selected_values() -> None:
                selected_values = {
                    str(item.data(Qt.ItemDataRole.UserRole))
                    for item in value_list.selectedItems()
                }
                if not selected_values or selected_values == set(unique_values):
                    self._proxy.clear_column_filter(column)
                else:
                    self._proxy.set_column_filter_values(column, selected_values)

            value_list.single_item_activated.connect(apply_single_value)
            apply_button.clicked.connect(apply_selected_values)
            filter_layout.addWidget(value_list)
            filter_layout.addWidget(apply_button)

            widget_action = QWidgetAction(menu)
            widget_action.setDefaultWidget(filter_widget)
            menu.addAction(widget_action)

            non_blank_values = {value for value in unique_values if value}
            if non_blank_values:
                menu.addSeparator()
                non_blank_action = menu.addAction("{Non-blanks}")
                non_blank_action.triggered.connect(
                    lambda: self._proxy.set_column_filter_values(column, non_blank_values)
                )

        menu.addSeparator()
        clear_filter_action = menu.addAction("Clear Filter")
        clear_filter_action.triggered.connect(
            lambda: self._proxy.clear_column_filter(column)
        )
        clear_all_action = menu.addAction("Clear All Filters")
        clear_all_action.triggered.connect(self._proxy.clear_all_column_filters)

        menu.addSeparator()
        sort_ascending_action = menu.addAction("Sort Ascending")
        sort_ascending_action.triggered.connect(
            lambda: self._table.sortByColumn(column, Qt.SortOrder.AscendingOrder)
        )
        sort_descending_action = menu.addAction("Sort Descending")
        sort_descending_action.triggered.connect(
            lambda: self._table.sortByColumn(column, Qt.SortOrder.DescendingOrder)
        )
        clear_sort_action = menu.addAction("Clear Sort")
        clear_sort_action.triggered.connect(
            lambda: header.setSortIndicator(-1, Qt.SortOrder.AscendingOrder)
        )
        menu.exec(header.mapToGlobal(pos))

    # ── Export: Excel ───────────────────────────────────────────────

    def export_to_excel(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Export to Excel", f"{self._title}.xlsx",
            "Excel (*.xlsx);;All Files (*)",
        )
        if not path:
            return
        try:
            with pd.ExcelWriter(path) as writer:
                self._model.df.to_excel(writer, index=False)
            if self._open_after.isChecked():
                import os; os.startfile(path)
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    # ── Export: Word ────────────────────────────────────────────────

    def export_to_word(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Export to Word", f"{self._title}.docx",
            "Word (*.docx);;All Files (*)",
        )
        if not path:
            return
        try:
            from docx import Document
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            doc = Document()
            nrows = self._model.rowCount()
            ncols = self._model.columnCount()
            table = doc.add_table(rows=nrows + 1, cols=ncols)
            table.style = "Table Grid"

            # Header row
            for j in range(ncols):
                cell = table.cell(0, j)
                cell.text = str(self._model.headerData(
                    j, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole
                ))
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="#244061"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

            # Data rows
            for row in range(nrows):
                for col in range(ncols):
                    idx = self._model.index(row, col)
                    text = self._model.data(idx, Qt.ItemDataRole.DisplayRole) or ""
                    bg = self._model.data(idx, Qt.ItemDataRole.BackgroundRole)
                    cell = table.cell(row + 1, col)
                    cell.text = str(text)
                    if bg and isinstance(bg, QColor):
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="{bg.name()}"/>'
                        )
                        cell._tc.get_or_add_tcPr().append(shading)

            doc.save(path)
            if self._open_after.isChecked():
                import os; os.startfile(path)
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    # ── Export: CSV ─────────────────────────────────────────────────

    def export_to_csv(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Export to CSV", f"{self._title}.csv",
            "CSV (*.csv);;All Files (*)",
        )
        if not path:
            return
        try:
            self._model.df.to_csv(path, index=False, encoding="utf-8-sig")
            if self._open_after.isChecked():
                import os; os.startfile(path)
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    # ── Copy to clipboard ───────────────────────────────────────────

    def copy_to_clipboard(self):
        """Copy entire table as tab-separated text."""
        lines = ["\t".join(str(c) for c in self._model.df.columns)]
        for r in range(self._model.rowCount()):
            row = []
            for c in range(self._model.columnCount()):
                idx = self._model.index(r, c)
                val = self._model.data(idx, Qt.ItemDataRole.DisplayRole)
                row.append(str(val) if val is not None else "")
            lines.append("\t".join(row))
        QApplication.clipboard().setText("\n".join(lines))

    # ── JSON save (matching original) ───────────────────────────────

    def save_to_json(self, filepath: str | Path):
        """Save table data + cell colors to JSON (matches FreeCAD format)."""
        data = []
        for col in range(self._model.columnCount()):
            text = self._model.headerData(
                col, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole
            )
            data.append({"row": 0, "col": col, "text": str(text), "color": ""})

        for row in range(self._model.rowCount()):
            for col in range(self._model.columnCount()):
                idx = self._model.index(row, col)
                text = self._model.data(idx, Qt.ItemDataRole.DisplayRole) or ""
                bg = self._model.data(idx, Qt.ItemDataRole.BackgroundRole)
                color = bg.name() if isinstance(bg, QColor) else ""
                data.append({
                    "row": row + 1, "col": col,
                    "text": str(text), "color": color,
                })

        Path(filepath).write_text(
            json.dumps(data, indent=4, ensure_ascii=False), encoding="utf-8"
        )


class ColumnValueFilterProxyModel(QtCore.QSortFilterProxyModel):
    """Filter rows by allowed display values on one or more columns."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._column_filters: dict[int, set[str]] = {}

    @staticmethod
    def _normalize_value(value) -> str:
        if value is None:
            return ""
        return str(value).strip()

    def set_column_filter_values(self, column: int, allowed_values: set[str] | list[str]) -> None:
        normalized = {
            self._normalize_value(value)
            for value in allowed_values
        }
        if normalized:
            self._column_filters[column] = normalized
        else:
            self._column_filters.pop(column, None)
        self.invalidateFilter()

    def get_column_filter_values(self, column: int) -> set[str]:
        return set(self._column_filters.get(column, set()))

    def clear_column_filter(self, column: int) -> None:
        if column in self._column_filters:
            del self._column_filters[column]
            self.invalidateFilter()

    def clear_all_column_filters(self) -> None:
        if not self._column_filters:
            return
        self._column_filters.clear()
        self.invalidateFilter()

    def filterAcceptsRow(self, source_row, source_parent):
        if not super().filterAcceptsRow(source_row, source_parent):
            return False
        model = self.sourceModel()
        if model is None:
            return True
        for column, allowed_values in self._column_filters.items():
            index = model.index(source_row, column, source_parent)
            value = self._normalize_value(
                model.data(index, Qt.ItemDataRole.DisplayRole)
            )
            if value not in allowed_values:
                return False
        return True

"""
Export Markdown help content to DOCX.

Uses ``python-docx`` to create a styled Word document from Markdown source.
Supports headings, paragraphs, tables, code blocks, blockquotes, and images.
"""

from __future__ import annotations

from pathlib import Path
from typing import Sequence

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

from markdown_it import MarkdownIt
from markdown_it.token import Token


# ═══════════════════════════════════════════════════════════════════════════
# Document setup
# ═══════════════════════════════════════════════════════════════════════════

def _create_docx(rtl: bool = False) -> docx.Document:
    """Create a new styled DOCX document."""
    doc = docx.Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "B Nazanin" if rtl else "Calibri"
    font.size = Pt(12 if rtl else 11)

    # Heading styles
    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = RGBColor(0x24, 0x40, 0x61)
        hs.font.name = "B Nazanin" if rtl else "Calibri"
        sizes = {1: 20, 2: 16, 3: 13}
        hs.font.size = Pt(sizes.get(i, 12))

    return doc


# ═══════════════════════════════════════════════════════════════════════════
# Markdown → DOCX pipeline
# ═══════════════════════════════════════════════════════════════════════════

def _parse_tokens(text: str) -> list[Token]:
    md = MarkdownIt("commonmark", {"html": True})
    md.enable(["table", "strikethrough"])
    return md.parse(text)


def _tokens_to_docx(doc: docx.Document, tokens: list[Token], rtl: bool = False):
    """Walk token stream and render to DOCX."""
    i = 0
    align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT

    while i < len(tokens):
        tok = tokens[i]

        if tok.type == "heading_open":
            level = int(tok.tag[1])
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                text = tokens[i + 1].content
                p = doc.add_heading(text, level=level)
                if rtl:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                i += 3
                continue

        elif tok.type == "paragraph_open":
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                text = tokens[i + 1].content
                p = doc.add_paragraph(text)
                p.alignment = align
                i += 3
                continue

        elif tok.type == "fence":
            # Code block
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(tok.content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # Light gray background via shading
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
            p._element.get_or_add_pPr().append(shading)
            i += 1
            continue

        elif tok.type == "blockquote_open":
            depth = 1
            j = i + 1
            bq_text = ""
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                elif tokens[j].type == "inline":
                    bq_text += tokens[j].content + "\n"
                j += 1
            p = doc.add_paragraph(bq_text.strip(), style="Quote")
            p.alignment = align
            i = j
            continue

        elif tok.type == "table_open":
            headers, rows = _extract_table(tokens, i)
            if headers:
                _add_table(doc, headers, rows)
            while i < len(tokens) and tokens[i].type != "table_close":
                i += 1
            i += 1
            continue

        i += 1


def _extract_table(tokens: list[Token], start: int):
    """Extract headers and rows from table token sequence."""
    headers = []
    rows = []
    current_row = []
    in_head = False
    in_body = False

    for i in range(start, len(tokens)):
        tok = tokens[i]
        if tok.type == "table_close":
            break
        if tok.type == "thead_open":
            in_head = True
        elif tok.type == "thead_close":
            in_head = False
        elif tok.type == "tbody_open":
            in_body = True
        elif tok.type == "tbody_close":
            in_body = False
        elif tok.type == "tr_open":
            current_row = []
        elif tok.type == "tr_close":
            if in_head:
                headers = current_row
            elif in_body:
                rows.append(current_row)
        elif tok.type == "inline":
            current_row.append(tok.content)

    return headers, rows


def _add_table(doc: docx.Document, headers: list[str], rows: list[list[str]]):
    """Add a formatted table to the document."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count, style="Table Grid")

    # Header row
    header_row = table.rows[0]
    for j, text in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = text
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
        if run:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="244061"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, text in enumerate(row_data):
            if j < col_count:
                row.cells[j].text = text

    doc.add_paragraph()  # spacing after table


# ═══════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════

def md_to_docx(
    md_files: Sequence[Path | str],
    output: Path | str,
    *,
    rtl: bool = False,
    title: str = "civilTools Help Manual",
):
    """Convert Markdown files to a single DOCX document.

    Parameters
    ----------
    md_files : sequence of paths
        Markdown files to include, in order.
    output : path
        Output DOCX file.
    rtl : bool
        Right-to-left layout.
    title : str
        Document title (used in heading).
    """
    doc = _create_docx(rtl=rtl)

    # Title
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for md_path in md_files:
        md_path = Path(md_path)
        text = md_path.read_text("utf-8")
        tokens = _parse_tokens(text)
        doc.add_page_break()
        _tokens_to_docx(doc, tokens, rtl=rtl)

    Path(output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return Path(output)

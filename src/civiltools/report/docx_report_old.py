"""
DOCX report generator for civilTools — uses python-docx + math2docx.

Generates structural engineering reports from BuildingModel data.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from civiltools.report.report_config import ReportConfig
from civiltools.report.strings import get_string


# ═══════════════════════════════════════════════════════════════════════════
# Document creation
# ═══════════════════════════════════════════════════════════════════════════

def _create_styled_doc(config: ReportConfig) -> Document:
    """Create a Document with custom styles."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = config.font_name if config.is_rtl else "Calibri"
    font.size = Pt(11)

    # Heading styles
    colors = {
        1: RGBColor(30, 60, 120),
        2: RGBColor(50, 100, 50),
        3: RGBColor(80, 80, 80),
    }
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}

    for level in (1, 2, 3):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = config.font_name if config.is_rtl else "Calibri"
        heading_style.font.size = sizes[level]
        heading_style.font.color.rgb = colors[level]
        heading_style.font.bold = True

    return doc


def add_table_of_contents(doc: Document):
    """Insert a TOC field (updated when user opens in Word)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = docx.oxml.OxmlElement("w:fldChar")
    fld_char_begin.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = docx.oxml.OxmlElement("w:instrText")
    instr.set(docx.oxml.ns.qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = docx.oxml.OxmlElement("w:fldChar")
    fld_char_end.set(docx.oxml.ns.qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    """Add a two-column key/value table to the document."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Shading Accent 1"
    for i, (key, val) in enumerate(rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(val)
    doc.add_paragraph()  # spacing


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a data table with headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.cell(r + 1, c).text = str(val)

    doc.add_paragraph()


def add_formula_section(doc: Document, steps: list[tuple[str, str]]):
    """Add LaTeX formulas to DOCX using math2docx."""
    try:
        from math2docx import Math2Docx
    except ImportError:
        # Fallback: plain text
        for desc, latex in steps:
            p = doc.add_paragraph()
            p.add_run(f"{desc}: ").bold = True
            p.add_run(latex)
        return

    m2d = Math2Docx()
    for desc, latex in steps:
        doc.add_paragraph(desc, style="Heading 3")
        try:
            m2d.add_math(doc, latex)
        except Exception:
            doc.add_paragraph(latex)


# ═══════════════════════════════════════════════════════════════════════════
# High-level report builder
# ═══════════════════════════════════════════════════════════════════════════

def create_docx_report(
    model,
    config: ReportConfig | None = None,
    output_path: str | Path | None = None,
) -> Path:
    """Generate a full DOCX report from a BuildingModel.

    Parameters
    ----------
    model : BuildingModel
        The building model with populated seismic_params.
    config : ReportConfig, optional
        Report configuration.
    output_path : str or Path, optional
        Output file path.
    """
    from civiltools.report.latex_str import full_earthquake_calculation

    config = config or ReportConfig()
    doc = _create_styled_doc(config)

    # Title page
    title = model.project_name or "Structural Report"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(30, 60, 120)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("civilTools Structural Engineering Report")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # Table of contents
    if config.include_table_of_contents:
        doc.add_heading(get_string("TABLE_OF_CONTENTS", config.language), level=1)
        add_table_of_contents(doc)
        doc.add_page_break()

    for section in config.active_sections:
        if section == "project_info":
            doc.add_heading(config.get_section_name(section), level=1)
            rows = [
                (get_string("PROJECT_INFO", config.language), model.project_name),
                (get_string("CITY", config.language), model.location),
                (get_string("NO_STORIES", config.language), str(len(model.stories))),
                (get_string("HEIGHT_METER", config.language), f"{model.total_height:.1f}"),
            ]
            add_key_value_table(doc, rows)

        elif section == "seismic_params":
            doc.add_heading(config.get_section_name(section), level=1)
            sp = model.seismic_params
            if sp:
                rows = [
                    (get_string("SOIL_TYPE", config.language), str(sp.get("soil_type", ""))),
                    (get_string("RISK_LEVEL", config.language), str(sp.get("zone", sp.get("risk_level", "")))),
                    (get_string("IMPORTANCE_FACTOR", config.language), str(sp.get("I", sp.get("importance_factor", "")))),
                    (get_string("DESIGN_BASE_ACCELERATION", config.language), str(sp.get("A", ""))),
                ]
                add_key_value_table(doc, rows)

        elif section == "earthquake_formulation":
            doc.add_heading(config.get_section_name(section), level=1)
            sp = model.seismic_params
            if sp:
                steps = full_earthquake_calculation(sp, "x")
                add_formula_section(doc, steps)

        elif section == "earthquake_values":
            doc.add_heading(config.get_section_name(section), level=1)
            sp = model.seismic_params
            if sp:
                steps = full_earthquake_calculation(sp, "y")
                add_formula_section(doc, steps)

        elif section in ("drift", "design_results", "json_tables"):
            doc.add_heading(config.get_section_name(section), level=1)
            doc.add_paragraph(
                f"[{section}] — Content will be populated from ETABS results."
            )

        else:
            doc.add_heading(config.get_section_name(section), level=1)
            doc.add_paragraph(f"[Section: {section} — content pending]")

    # Save
    if output_path is None:
        name = model.project_name.replace(" ", "_") or "report"
        output_path = Path(f"{name}_report.docx")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path

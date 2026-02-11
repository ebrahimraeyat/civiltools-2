"""
Earthquake factor export to Word — ported from civilTools/exporter/export_to_word.py.

Generates a Persian (RTL) DOCX with project specifications,
structural system properties, and seismic coefficients.
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_data_from_model(building):
    """Extract project, structural system, and result data from a Building."""
    X = building.x_system
    Y = building.y_system

    prop = {
        '': '',
        'محل اجرای پروژه': building.city,
        'کاربری ساختمان': 'مسکونی',
        'ضریب اهمیت': building.importance_factor,
        'تعداد طبقات': building.number_of_story,
        'ارتفاع ساختمان  )متر(': building.height,
        'سطح خطر نسبی': building.risk_level,
        'شتاب مبنای طرح': building.acc,
        'نوع خاک': building.soil_type,
    }

    if X == Y:
        prop['سیستم سازه ای در دو راستا'] = X.lateral_type
    else:
        prop['سیستم سازه ای در راستای x'] = X.lateral_type
        prop['سیستم سازه ای در راستای y'] = Y.lateral_type

    struc = {
        '': ('راستای x', 'راستای y'),
        'سیستم سازه': (X.lateral_type, Y.lateral_type),
        'ضریب رفتار': (X.Ru, Y.Ru),
        'ضریب اضافه مقاومت': (X.phi0, Y.phi0),
        'ضریب بزرگنمایی جابجایی': (X.cd, Y.cd),
        'ارتفاع مجاز  )متر(': (X.max_height, Y.max_height),
    }

    result = {
        'زمان تناوب تجربی': (building.tx_exp, building.ty_exp),
        'زمان تناوب تحلیلی': (building.tx_an, building.ty_an),
        'ضریب بازتاب': (building.bx, building.by),
        'ضریب زلزله': (building.results[1], building.results[2]),
        'ضریب توزیع در ارتفاع': (building.kx, building.ky),
        'ضریب زلزله دریفت': (building.results_drift[1], building.results_drift[2]),
        'ضریب توزیع در ارتفاع دریفت': (building.kx_drift, building.ky_drift),
    }

    return prop, struc, result


def export(building=None, filename=None, doc=None):
    """Export earthquake factor data for *building* to a Word document.

    Parameters
    ----------
    building : Building
        The computed building object (with results already calculated).
    filename : str | None
        Output .docx path. If given the document is saved to disk.
    doc : Document | None
        An existing python-docx Document to append to.
        If *None* a new blank document is created.

    Returns
    -------
    Document
        The python-docx Document (useful when chaining multiple exports).
    """
    prop, struc, result = get_data_from_model(building)

    if doc is None:
        doc = Document()

    table_style = 'Table Grid'
    # Try to use a nicer built-in style if available
    for preferred in ('List Table 4 Accent 5', 'Light Shading Accent 1'):
        if preferred in [s.name for s in doc.styles]:
            table_style = preferred
            break

    # ── Project specifications ──────────────────────────────────────
    h = doc.add_heading('محاسبه ضریب زلزله', level=0)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    h = doc.add_heading('مشخصات پروژه', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table_prop = doc.add_table(rows=0, cols=2, style=doc.styles[table_style])
    try:
        table_prop.direction = WD_TABLE_DIRECTION.RTL
    except Exception:
        pass
    for key, value in prop.items():
        row_cells = table_prop.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    # ── Structural system ───────────────────────────────────────────
    h = doc.add_heading('مشخصات سیستم سازه ای', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    struc_table = doc.add_table(rows=0, cols=3, style=doc.styles[table_style])
    for key, value in struc.items():
        row_cells = struc_table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value[0])
        row_cells[2].text = str(value[1])

    # ── Earthquake coefficients ─────────────────────────────────────
    h = doc.add_heading('ضریب زلزله', level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    result_table = doc.add_table(rows=1, cols=3, style=doc.styles[table_style])
    hdr_cells = result_table.rows[0].cells
    hdr_cells[1].text = 'راستای x'
    hdr_cells[2].text = 'راستای y'

    for key, value in result.items():
        row_cells = result_table.add_row().cells
        row_cells[0].text = key
        try:
            row_cells[1].text = f'{value[0]:.3f}'
        except (ValueError, TypeError):
            row_cells[1].text = str(value[0])
        try:
            row_cells[2].text = f'{value[1]:.3f}'
        except (ValueError, TypeError):
            row_cells[2].text = str(value[1])

    if filename:
        doc.save(filename)

    return doc

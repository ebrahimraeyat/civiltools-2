"""
DOCX report generator for civilTools — full English structural report.

Uses python-docx + math2docx.  Accepts a ``ReportData`` instance
produced by ``data_extractor.extract_report_data`` and renders every
active section to a styled Word document.
"""

from __future__ import annotations

import io
import re
import tempfile
from pathlib import Path
from typing import Any

import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from civiltools.report.report_config import ReportConfig
from civiltools.report.strings import get_string, ASCE7_IRREGULARITY_DESC
from civiltools.report.data_extractor import ReportData


# ── Farsi → English translations ────────────────────────────────────────
_RISK_EN = {
    "\u06a9\u0645": "Low",
    "\u0645\u062a\u0648\u0633\u0637": "Moderate",
    "\u0632\u06cc\u0627\u062f": "High",
    "\u062e\u06cc\u0644\u06cc \u0632\u06cc\u0627\u062f": "Very High",
}

def _en_risk(risk_level) -> str:
    """Translate Farsi risk level to English."""
    return _RISK_EN.get(str(risk_level), str(risk_level))

def _en_soil(soil_type) -> str:
    """Format soil type for display."""
    return f"Type {soil_type}" if soil_type else ""


# ═══════════════════════════════════════════════════════════════════════════
# Document creation
# ═══════════════════════════════════════════════════════════════════════════

def _create_styled_doc(config: ReportConfig) -> Document:
    """Create a Document with custom styles."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    colors = {
        1: RGBColor(30, 60, 120),
        2: RGBColor(50, 100, 50),
        3: RGBColor(80, 80, 80),
    }
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}

    for level in (1, 2, 3):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.size = sizes[level]
        heading_style.font.color.rgb = colors[level]
        heading_style.font.bold = True

    return doc


def _add_table_of_contents(doc: Document):
    """Insert a TOC field (updated when user opens in Word)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = docx.oxml.OxmlElement("w:fldChar")
    fld_begin.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    run2 = paragraph.add_run()
    instr = docx.oxml.OxmlElement("w:instrText")
    instr.set(docx.oxml.ns.qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_end = docx.oxml.OxmlElement("w:fldChar")
    fld_end.set(docx.oxml.ns.qn("w:fldCharType"), "end")
    run3._element.append(fld_end)


# ═══════════════════════════════════════════════════════════════════════════
# Table helpers
# ═══════════════════════════════════════════════════════════════════════════

def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    """Two-column key/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Shading Accent 1"
    for i, (key, val) in enumerate(rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(val)
    doc.add_paragraph()


def _add_data_table(doc: Document, headers: list[str], rows: list[list]):
    """Generic data table with bold header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.cell(r_idx + 1, c_idx).text = str(val)
    _format_table_font(table, body_size=9.0, header_size=9.5)
    doc.add_paragraph()
    return table


def _add_df_table(doc: Document, df, max_rows: int = 500):
    """Add a pandas DataFrame as a table."""
    headers = list(df.columns)
    rows = df.head(max_rows).values.tolist()
    _add_data_table(doc, headers, rows)


def _add_image(doc: Document, img_bytes: bytes, width_inches: float = 6.0):
    """Add PNG image from bytes."""
    stream = io.BytesIO(img_bytes)
    doc.add_picture(stream, width=Inches(width_inches))
    doc.add_paragraph()


def _clean_latex_cell(s: str) -> str:
    """Strip simple LaTeX markup for use in a plain text table cell."""
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)   # \text{...} → ...
    s = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', s)  # \mathrm{...} → ...
    s = re.sub(r'_\{([^}]+)\}', r'_\1', s)        # _{0} → _0
    s = re.sub(r'\\[a-zA-Z]+', '', s)              # remaining \cmds
    s = s.replace('{', '').replace('}', '')
    return s.strip()


def _try_render_array_as_table(doc: Document, latex: str) -> bool:
    """If *latex* contains a LaTeX array/tabular, render it as a Word table.

    Returns True on success so the caller can skip the image renderer.
    """
    m = re.search(
        r'\\begin\{(?:array|tabular)\}\{[^}]*\}(.*?)\\end\{(?:array|tabular)\}',
        latex, re.DOTALL,
    )
    if not m:
        return False

    content = m.group(1)
    # Rows are separated by \\
    row_strs = re.split(r'\\\\', content)
    rows: list[list[str]] = []
    for rs in row_strs:
        rs = rs.replace(r'\hline', '').strip()
        if not rs:
            continue
        cells = [_clean_latex_cell(c) for c in rs.split('&')]
        rows.append(cells)

    if len(rows) < 2:
        return False

    try:
        _add_data_table(doc, rows[0], rows[1:])
        return True
    except Exception:
        return False


def _add_formula_section(doc: Document, steps: list[tuple[str, str]]):
    """Add LaTeX formulas rendered as images, with plain-text fallback.

    Steps containing a LaTeX array/tabular environment are rendered as
    proper Word tables instead of images (matplotlib cannot render
    ``\\begin{array}`` in inline-math mode).
    """
    for desc, latex in steps:
        doc.add_paragraph(desc, style="Heading 3")
        # Prefer Word-table rendering for array environments
        if r'\begin{array}' in latex or r'\begin{tabular}' in latex:
            if _try_render_array_as_table(doc, latex):
                continue
        try:
            img_bytes = _render_latex_to_png(latex)
            stream = io.BytesIO(img_bytes)
            doc.add_picture(stream, width=Inches(5.5))
        except Exception:
            # Fallback: plain text
            p = doc.add_paragraph()
            p.add_run(latex).font.name = "Consolas"


def _latex_to_inline_text(latex: str) -> str:
    """Convert simple LaTeX expressions to compact inline text."""
    text = latex.strip()

    if r"\begin{array}" in text or r"\begin{tabular}" in text:
        text = re.sub(r"\\begin\{(?:array|tabular)\}\{[^}]*\}", "", text)
        text = re.sub(r"\\end\{(?:array|tabular)\}", "", text)
        text = text.replace(r"\hline", "")
        text = re.sub(r"\\\\", " ; ", text)
        text = text.replace("&", ", ")

    while True:
        updated = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
        if updated == text:
            break
        text = updated

    replacements = {
        r"\times": "×",
        r"\cdot": "×",
        r"\geq": "≥",
        r"\leq": "≤",
        r"\Rightarrow": "⇒",
        r"\checkmark": "✓",
        r"\quad": " ",
        r"\,": " ",
        r"\;": " ",
        r"\!": "",
        r"\left": "",
        r"\right": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"_\{([^}]*)\}", r"_\1", text)
    text = re.sub(r"\^\{([^}]*)\}", r"^\1", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\\[A-Za-z]+", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _format_table_font(table, body_size: float = 8.0, header_size: float = 8.5):
    """Apply compact font/spacing to a Word table."""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(header_size if row_idx == 0 else body_size)
                    if row_idx == 0:
                        run.bold = True


def _add_formula_summary_table(doc: Document, rows: list[list[str]]):
    """Add a compact 3-column table for earthquake calculations."""
    headers = ["Parameter", "Formula", "Substitution / Result"]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Shading Accent 1"
    table.autofit = False

    for col_idx, header in enumerate(headers):
        table.cell(0, col_idx).text = header

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = str(value)

    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(2.05)
    table.columns[2].width = Inches(2.80)
    _format_table_font(table, body_size=8.0, header_size=8.5)
    doc.add_paragraph()


def _build_earthquake_formula_rows(params: dict[str, Any], direction: str) -> list[list[str]]:
    """Build compact earthquake-calculation rows for a 3-column table."""
    try:
        from civiltools.report.latex_str import (
            b1_with_values,
            b_with_values,
            c_check_with_values,
            c_min_with_values,
            design_period_with_values,
            earthquake_c_with_values,
            k_with_values,
            n_with_values,
            period_with_values,
        )
    except ImportError:
        return []

    d = direction.lower()
    A = params.get("A", 0.3)
    I_ = params.get("I", 1.0)
    R = params.get(f"R{d}", params.get("R", 7.0))
    T_emp = params.get(f"T{d}", 0.5)
    T_an = params.get(f"T{d}_an", 0.6)
    T_design = params.get(f"T{d}_design", T_emp)
    B1 = params.get(f"B1{d}", 2.5)
    N = params.get(f"N{d}", 1.0)
    B = params.get(f"B{d}", 2.5)
    K = params.get(f"K{d}", 1.0)
    C = params.get(f"C{d}", 0.1)
    alpha = params.get("alpha", params.get("Ct", 0.07))
    beta = params.get("beta", 0.75)
    H = params.get("H", params.get("height", 10.0))
    soil_type = params.get("soil_type", "III")
    T0 = params.get("T0", 0.15)
    Ts = params.get("Ts", 0.70)
    S = params.get("S", 1.75)
    S0 = params.get("S0", 1.75)
    is_high = params.get("risk_level", 3) >= 3
    is_infill = params.get("is_infill", False)

    C_min = 0.12 * A * I_
    C_final = max(C, C_min)
    n_formula = (
        "N = 1 / (0.7T/Ts)^0.4 / 1.3"
        if is_high else
        "N = 1 / (0.7T/Ts)^0.5 / 1.45"
    )

    return [
        [
            "Empirical period T_emp",
            "T_emp = 0.8×α×H^β" if is_infill else "T_emp = α×H^β",
            _latex_to_inline_text(period_with_values(alpha, beta, H, T_emp, is_infill)),
        ],
        [
            "Design period T",
            "T = max(T_emp, min(T_an, 1.25×T_emp))",
            _latex_to_inline_text(design_period_with_values(T_emp, T_an, T_design)),
        ],
        [
            "Soil parameters",
            "Lookup from soil type table",
            f"Soil {soil_type}: T0={T0}, Ts={Ts}, S={S}, S0={S0}",
        ],
        [
            "Reflection coefficient B1",
            "B1 = piecewise(T, T0, Ts, S, S0)",
            _latex_to_inline_text(b1_with_values(T_design, T0, Ts, S, S0, B1)),
        ],
        [
            "N coefficient",
            n_formula,
            _latex_to_inline_text(n_with_values(T_design, Ts, N, is_high)),
        ],
        [
            "Reflection factor B",
            "B = B1×N",
            _latex_to_inline_text(b_with_values(B1, N, B)),
        ],
        [
            "Distribution exponent K",
            "K = 1 / (0.5T + 0.75) / 2",
            _latex_to_inline_text(k_with_values(T_design, K)),
        ],
        [
            "Base coefficient C",
            "C = A×B×I / R",
            _latex_to_inline_text(earthquake_c_with_values(A, B, I_, R, C)),
        ],
        [
            "Minimum coefficient C_min",
            "C_min = 0.12×A×I",
            _latex_to_inline_text(c_min_with_values(A, I_, C_min)),
        ],
        [
            "Final design coefficient",
            "C_design = max(C, C_min)",
            _latex_to_inline_text(c_check_with_values(C, C_min, C_final)),
        ],
    ]


def _render_latex_to_png(latex_str: str, dpi: int = 150, fontsize: int = 14) -> bytes:
    """Render a LaTeX string to PNG bytes using matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(8, 1.2))
    ax.axis("off")
    ax.text(0.5, 0.5, f"${latex_str}$",
            fontsize=fontsize, ha="center", va="center",
            transform=ax.transAxes)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                transparent=True, pad_inches=0.1)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════
# Section renderers
# ═══════════════════════════════════════════════════════════════════════════

def _section_model_settings(doc: Document, data: ReportData, lang: str):
    """Render model settings JSON as well-categorized tables."""
    doc.add_heading("Model Settings", level=1)
    ms = data.model_settings
    if not ms:
        doc.add_paragraph(
            "Model settings JSON not found. "
            "Please ensure that the model has been opened and saved at least once "
            "from the civilTools interface so that the settings file is generated "
            "in the table_results folder next to the ETABS model."
        )
        if data.building:
            doc.add_heading("Available Data from Building Object", level=2)
            b = data.building
            rows = [
                ("Soil Type", _en_soil(getattr(b, "soil_type", ""))),
                ("Risk Level", _en_risk(getattr(b, "risk_level", ""))),
                ("Importance Factor", str(getattr(b, "importance_factor", ""))),
                ("Height (m)", str(getattr(b, "height", ""))),
            ]
            _add_kv_table(doc, [(k, v) for k, v in rows if v])
        return

    # ── 1. Project Information ────────────────────────────────────────
    doc.add_heading("Project Information", level=2)
    info_rows = [
        ("Province", ms.get("ostan", "")),
        ("City", ms.get("city", "")),
        ("Relative Risk Level", _en_risk(ms.get("risk_level", ""))),
        ("Soil Type", _en_soil(ms.get("soil_type", ""))),
        ("Importance Factor", ms.get("importance_factor", "")),
    ]
    _add_kv_table(doc, [(k, str(v)) for k, v in info_rows if v])

    # ── 2. Gravity Loads ──────────────────────────────────────────────
    doc.add_heading("Gravity Loads", level=2)
    _load_map = [
        ("Dead Load (DL)", "dead_combobox"),
        ("Super Dead Load (SD)", "sdead_combobox"),
        ("Partition Dead", "partition_dead_combobox"),
        ("Live Load", "live_combobox"),
        ("Reducible Live", "lred_combobox"),
        ("Parking Live", "live_parking_combobox"),
        ("Partition Live", "partition_live_combobox"),
        ("Live Load 0.5", "live5_combobox"),
        ("Reducible Live 0.5", "lred5_combobox"),
        ("Roof Live", "lroof_combobox"),
        ("Mass", "mass_combobox"),
        ("Vertical Seismic (EV)", "ev_combobox"),
        ("Modal", "modal_combobox"),
    ]
    load_rows = [(label, ms.get(key, "")) for label, key in _load_map]
    _add_kv_table(doc, [(k, str(v)) for k, v in load_rows if v])

    # Retaining wall loads
    if ms.get("retaining_wall_groupbox"):
        doc.add_heading("Retaining Wall Loads", level=3)
        rw_rows = [
            ("Lateral Soil Pressure X+", ms.get("hxp_combobox", "")),
            ("Lateral Soil Pressure X-", ms.get("hxn_combobox", "")),
            ("Lateral Soil Pressure Y+", ms.get("hyp_combobox", "")),
            ("Lateral Soil Pressure Y-", ms.get("hyn_combobox", "")),
        ]
        _add_kv_table(doc, [(k, str(v)) for k, v in rw_rows if v])

    # ── 3. Structural System (Primary) ────────────────────────────────
    is_dual = bool(ms.get("activate_second_system"))
    _sys_label = "Structural System" if not is_dual else "Structural System (Lower)"
    doc.add_heading(_sys_label, level=2)
    sys_rows = [
        ("X System", ms.get("x_system_name", "")),
        ("X Lateral System", ms.get("x_lateral_name", "")),
        ("Ru (X)", str(ms.get("Rux", ""))),
        ("Cd (X)", str(ms.get("cdx", ""))),
        ("Y System", ms.get("y_system_name", "")),
        ("Y Lateral System", ms.get("y_lateral_name", "")),
        ("Ru (Y)", str(ms.get("Ruy", ""))),
        ("Cd (Y)", str(ms.get("cdy", ""))),
        ("Bottom Story", ms.get("bot_x_combo", "")),
        ("Top Story", ms.get("top_x_combo", "")),
        *(([("Top Story for Height", ms.get("top_story_for_height", ""))]) if not is_dual else []),
        ("Number of Stories", str(ms.get("no_of_story_x", "") or ms.get("no_of_story", ""))),
        ("Building Height (m)", str(ms.get("height_x", ""))),
        ("Infill Panel", "Yes" if ms.get("infill") else "No"),
    ]
    _add_kv_table(doc, [(k, v) for k, v in sys_rows if v and v != "None"])

    # ── 4. Structural System (Secondary) ──────────────────────────────
    if is_dual:
        doc.add_heading("Structural System (Upper)", level=2)
        sys2_rows = [
            ("X System", ms.get("x_system_name_1", "")),
            ("X Lateral System", ms.get("x_lateral_name_1", "")),
            ("Ru (X)", str(ms.get("Rux1", ""))),
            ("Cd (X)", str(ms.get("cdx1", ""))),
            ("Y System", ms.get("y_system_name_1", "")),
            ("Y Lateral System", ms.get("y_lateral_name_1", "")),
            ("Ru (Y)", str(ms.get("Ruy1", ""))),
            ("Cd (Y)", str(ms.get("cdy1", ""))),
            ("Bottom Story", ms.get("bot_x1_combo", "")),
            ("Top Story", ms.get("top_x1_combo", "")),
            ("Top Story for Height", ms.get("top_story_for_height1", "")),
            ("Number of Stories", str(ms.get("no_of_story_x1", ""))),
            ("Building Height (m)", str(ms.get("height_x1", ""))),
            ("Infill Panel", "Yes" if ms.get("infill_1") else "No"),
            ("Stiffness 10x", "Yes" if ms.get("special_case") else "No"),
        ]
        _add_kv_table(doc, [(k, v) for k, v in sys2_rows if v and v != "None"])

    # ── 5. Static Earthquake Load Cases ───────────────────────────────
    doc.add_heading("Static Earthquake Load Cases", level=2)
    eq_header = ["Load Case", "X Direction", "Y Direction"]
    eq_rows = [
        ["EQ without Eccentricity", ms.get("ex_combobox", ""), ms.get("ey_combobox", "")],
        ["EQ + Positive Eccentricity", ms.get("exp_combobox", ""), ms.get("eyp_combobox", "")],
        ["EQ + Negative Eccentricity", ms.get("exn_combobox", ""), ms.get("eyn_combobox", "")],
        ["EQ Drift (no eccentricity)", ms.get("ex_drift_combobox", ""), ms.get("ey_drift_combobox", "")],
        ["EQ Drift + Positive Ecc.", ms.get("exp_drift_combobox", ""), ms.get("eyp_drift_combobox", "")],
        ["EQ Drift + Negative Ecc.", ms.get("exn_drift_combobox", ""), ms.get("eyn_drift_combobox", "")],
    ]
    eq_rows = [r for r in eq_rows if r[1] or r[2]]
    if eq_rows:
        _add_data_table(doc, eq_header, eq_rows)

    rho_rows = [
        ("Redundancy Factor (rho) X", ms.get("rhox_combobox", "")),
        ("Redundancy Factor (rho) Y", ms.get("rhoy_combobox", "")),
    ]
    _add_kv_table(doc, [(k, str(v)) for k, v in rho_rows if v])

    # Second system earthquake load cases
    if is_dual:
        doc.add_heading("Static EQ Load Cases (Upper System)", level=3)
        eq2_rows = [
            ["EQ without Eccentricity", ms.get("ex1_combobox", ""), ms.get("ey1_combobox", "")],
            ["EQ + Positive Eccentricity", ms.get("exp1_combobox", ""), ms.get("eyp1_combobox", "")],
            ["EQ + Negative Eccentricity", ms.get("exn1_combobox", ""), ms.get("eyn1_combobox", "")],
            ["EQ Drift (no eccentricity)", ms.get("ex1_drift_combobox", ""), ms.get("ey1_drift_combobox", "")],
            ["EQ Drift + Positive Ecc.", ms.get("exp1_drift_combobox", ""), ms.get("eyp1_drift_combobox", "")],
            ["EQ Drift + Negative Ecc.", ms.get("exn1_drift_combobox", ""), ms.get("eyn1_drift_combobox", "")],
        ]
        eq2_rows = [r for r in eq2_rows if r[1] or r[2]]
        if eq2_rows:
            _add_data_table(doc, eq_header, eq2_rows)

    # ── 6. Dynamic Analysis ───────────────────────────────────────────
    if ms.get("dynamic_analysis_groupbox"):
        doc.add_heading("Dynamic Analysis", level=2)
        doc.add_paragraph(
            f"X Scale Factor: {ms.get('x_scalefactor_combobox', '')}, "
            f"Y Scale Factor: {ms.get('y_scalefactor_combobox', '')}"
        )
        if ms.get("combination_response_spectrum_checkbox"):
            dyn_header = ["Load Case", "X Direction", "Y Direction"]
            dyn_rows = [
                ["Spectral (no ecc.)", ms.get("sx_combobox", ""), ms.get("sy_combobox", "")],
                ["Spectral (with ecc.)", ms.get("sxe_combobox", ""), ms.get("sye_combobox", "")],
                ["Spectral Drift (no ecc.)", ms.get("sx_drift_combobox", ""), ms.get("sy_drift_combobox", "")],
                ["Spectral Drift (with ecc.)", ms.get("sxe_drift_combobox", ""), ms.get("sye_drift_combobox", "")],
            ]
            dyn_rows = [r for r in dyn_rows if r[1] or r[2]]
            if dyn_rows:
                _add_data_table(doc, dyn_header, dyn_rows)

    # ── 7. Analytical Periods ─────────────────────────────────────────
    doc.add_heading("Analytical Periods", level=2)
    period_rows = [
        ("T analytical X (s)", str(ms.get("tx_an", ""))),
        ("T analytical Y (s)", str(ms.get("ty_an", ""))),
    ]
    if is_dual:
        period_rows.extend([
            ("T analytical X - upper (s)", str(ms.get("tx1_an", ""))),
            ("T analytical Y - upper (s)", str(ms.get("ty1_an", ""))),
        ])
    _add_kv_table(doc, [(k, v) for k, v in period_rows if v and v != "None"])

    # ── 8. Structural Material ────────────────────────────────────────
    material = "Concrete" if ms.get("concrete_radiobutton") else "Steel"
    doc.add_heading("Structural Material", level=2)
    _add_kv_table(doc, [("Material", material)])

    # ── 9. Irregularities ─────────────────────────────────────────────
    doc.add_heading("Irregularities", level=2)

    doc.add_heading("Plan Irregularities", level=3)
    plan_items = []
    if ms.get("torsional_irregularity_groupbox"):
        if ms.get("extreme_torsion_irregular_checkbox"):
            plan_items.append(("Extreme Torsional Irregularity", True))
        elif ms.get("torsion_irregular_checkbox"):
            plan_items.append(("Torsional Irregularity", True))
    else:
        plan_items.append(("Torsional Irregularity", False))
    plan_items.extend([
        ("Re-entrant Corner", ms.get("reentrance_corner_checkbox", False)),
        ("Diaphragm Discontinuity", ms.get("diaphragm_discontinuity_checkbox", False)),
        ("Out-of-Plane Offset", ms.get("out_of_plane_offset_checkbox", False)),
        ("Non-Parallel System", ms.get("nonparallel_system_checkbox", False)),
    ])
    _add_checkbox_table(doc, plan_items)

    doc.add_heading("Vertical Irregularities", level=3)
    vert_items = []
    if ms.get("stiffness_soft_story_groupbox"):
        if ms.get("extreme_stiffness_irregular_checkbox"):
            vert_items.append(("Extreme Soft Story", True))
        elif ms.get("stiffness_irregular_checkbox"):
            vert_items.append(("Soft Story", True))
    else:
        vert_items.append(("Stiffness Irregularity", False))
    vert_items.extend([
        ("Weight (Mass) Irregularity", ms.get("weight_mass_checkbox", False)),
        ("Geometric Irregularity", ms.get("geometric_checkbox", False)),
        ("In-Plane Discontinuity", ms.get("in_plane_discontinuity_checkbox", False)),
    ])
    if ms.get("lateral_strength_weak_story_groupbox"):
        if ms.get("extreme_strength_irregular_checkbox"):
            vert_items.append(("Extreme Weak Story", True))
        elif ms.get("strength_irregular_checkbox"):
            vert_items.append(("Weak Story", True))
    else:
        vert_items.append(("Lateral Strength Irregularity", False))
    _add_checkbox_table(doc, vert_items)


def _add_checkbox_table(doc: Document, items: list[tuple[str, bool]]):
    """Render a checkbox list with ASCE 7-16 descriptions for detected irregularities."""
    table = doc.add_table(rows=len(items), cols=1)
    table.style = "Light Shading Accent 1"
    for i, (label, checked) in enumerate(items):
        mark = "\u2713" if checked else "\u2610"
        table.cell(i, 0).text = f"{mark}  {label}"
    doc.add_paragraph()

    # Add ASCE 7-16 penalty descriptions for each detected irregularity
    detected = [label for label, checked in items if checked]
    if detected:
        doc.add_heading("ASCE 7-16 Irregularity Descriptions", level=4)
        for label in detected:
            desc = ASCE7_IRREGULARITY_DESC.get(label)
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(desc)
        doc.add_paragraph()


def _section_project_info(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("PROJECT_INFO", lang), level=1)
    b = data.building
    # Prefer reading story count from model settings (avoids ETABS off-by-one)
    ms = data.model_settings or {}
    no_stories = (
        ms.get("no_of_story_x")
        or ms.get("no_of_story")
        or len(data.stories)
    )
    rows = [
        (get_string("PROJECT_INFO", lang), data.project_name),
        (get_string("CITY", lang), data.location),
        (get_string("NO_STORIES", lang), str(no_stories)),
        (get_string("HEIGHT_METER", lang), f"{data.total_height:.1f}"),
    ]
    if b:
        rows.extend([
            (get_string("SOIL_TYPE", lang), _en_soil(getattr(b, "soil_type", ""))),
            (get_string("RISK_LEVEL", lang), _en_risk(getattr(b, "risk_level", ""))),
            (get_string("IMPORTANCE_FACTOR", lang), str(getattr(b, "importance_factor", ""))),
            (get_string("DESIGN_BASE_ACCELERATION", lang), str(getattr(b, "acc", ""))),
        ])
    _add_kv_table(doc, rows)


def _section_structural_system(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("STRUCTURAL_SYSTEM", lang), level=1)
    b = data.building
    if not b:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    for direction, sys_attr in [("X", "x_system"), ("Y", "y_system")]:
        sys = getattr(b, sys_attr, None)
        if sys is None:
            continue
        doc.add_heading(f"{get_string(direction + '_DIRECTION', lang)}", level=2)
        rows = [
            (get_string("LATERAL_TYPE", lang), str(getattr(sys, "lateral_type", getattr(sys, "system_name", str(sys))))),
            (get_string("RU_FACTOR", lang), str(getattr(sys, "Ru", ""))),
            (get_string("PHI0_FACTOR", lang), str(getattr(sys, "phi0", ""))),
            (get_string("CD_FACTOR", lang), str(getattr(sys, "cd", ""))),
            (get_string("MAX_HEIGHT", lang), str(getattr(sys, "max_height", ""))),
        ]
        _add_kv_table(doc, rows)


def _section_seismic_params(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("SEISMIC_PARAMS", lang), level=1)
    b = data.building
    if not b:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return
    rows = [
        (get_string("SOIL_TYPE", lang), _en_soil(getattr(b, "soil_type", ""))),
        (get_string("RISK_LEVEL", lang), _en_risk(getattr(b, "risk_level", ""))),
        (get_string("DESIGN_BASE_ACCELERATION", lang), str(getattr(b, "acc", ""))),
        (get_string("IMPORTANCE_FACTOR", lang), str(getattr(b, "importance_factor", ""))),
        (get_string("INFILL_PANNEL", lang), get_string("YES" if getattr(b, "is_infill", False) else "NO", lang)),
        (get_string("NO_STORIES", lang), str(getattr(b, "number_of_story", len(data.stories)))),
        (get_string("HEIGHT_METER", lang), f"{getattr(b, 'height', data.total_height):.1f}"),
    ]
    _add_kv_table(doc, rows)

    # Period summary
    doc.add_heading("Period Summary", level=2)
    period_rows = []
    for d_label, attr_exp, attr_an, attr_design in [
        ("X", "tx_exp", "tx_an", "tx"),
        ("Y", "ty_exp", "ty_an", "ty"),
    ]:
        t_exp = getattr(b, attr_exp, None)
        t_an = getattr(b, attr_an, None)
        t_des = getattr(b, attr_design, None)
        period_rows.append([
            d_label,
            f"{t_exp:.3f}" if t_exp else "—",
            f"{t_an:.3f}" if t_an else "—",
            f"{t_des:.3f}" if t_des else "—",
        ])
    _add_data_table(doc,
        ["Direction", "T empirical (s)", "T analytical (s)", "T design (s)"],
        period_rows,
    )


def _section_earthquake_formulation(doc: Document, data: ReportData, lang: str):
    """Earthquake coefficient formulation in compact 3-column tables."""
    doc.add_heading(get_string("EARTHQUAKE_COEFFICIENT_CALC", lang), level=1)
    b = data.building
    if not b:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    is_dual = b.building2 is not None
    sys_label = "Structural System (Lower)" if is_dual else "Structural System"

    # ── Lower / only system ───────────────────────────────────────
    doc.add_heading(sys_label, level=2)
    for direction in ("x", "y"):
        try:
            sp = _building_to_seismic_dict(b, direction)
            if sp:
                doc.add_heading(
                    get_string(direction.upper() + "_DIRECTION", lang), level=3
                )
                rows = _build_earthquake_formula_rows(sp, direction)
                if rows:
                    _add_formula_summary_table(doc, rows)
        except Exception:
            pass

    # ── Upper system (dual system only) ────────────────────────────
    if is_dual:
        doc.add_heading("Structural System (Upper)", level=2)
        for direction in ("x", "y"):
            try:
                sp2 = _building_to_seismic_dict(b.building2, direction)
                if sp2:
                    doc.add_heading(
                        get_string(direction.upper() + "_DIRECTION", lang), level=3
                    )
                    rows = _build_earthquake_formula_rows(sp2, direction)
                    if rows:
                        _add_formula_summary_table(doc, rows)
            except Exception:
                pass

        # ── Combined system (using combined period) ──────────────────
        doc.add_heading("Combined System (Overall Height)", level=2)
        for direction in ("x", "y"):
            try:
                sp_all = _building_all_to_seismic_dict(b, direction)
                if sp_all:
                    doc.add_heading(
                        get_string(direction.upper() + "_DIRECTION", lang), level=3
                    )
                    rows = _build_earthquake_formula_rows(sp_all, direction)
                    if rows:
                        _add_formula_summary_table(doc, rows)
            except Exception:
                pass


def _section_earthquake_values(doc: Document, data: ReportData, lang: str):
    """Coefficient summary table for all systems and directions."""
    doc.add_heading("Earthquake Coefficient Summary", level=1)
    b = data.building
    if not b:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    is_dual = b.building2 is not None

    headers = ["System", "Direction", "C (Design Coefficient)"]
    rows = []

    def _cx_cy(results):
        if results and results[0] is True:
            _, cx, cy = results
            return cx, cy
        return None, None

    if is_dual:
        # Lower system — individual period
        cx, cy = _cx_cy(getattr(b, "results", None))
        if cx: rows.append(["Lower", "X", f"{cx:.4f}"])
        if cy: rows.append(["Lower", "Y", f"{cy:.4f}"])
        # Combined period — lower portion
        cx, cy = _cx_cy(getattr(b, "results_all_bot", None))
        if cx: rows.append(["Lower (combined T)", "X", f"{cx:.4f}"])
        if cy: rows.append(["Lower (combined T)", "Y", f"{cy:.4f}"])
        # Combined period — upper portion
        cx, cy = _cx_cy(getattr(b, "results_all_top", None))
        if cx: rows.append(["Upper (combined T)", "X", f"{cx:.4f}"])
        if cy: rows.append(["Upper (combined T)", "Y", f"{cy:.4f}"])
        # Upper system — individual period
        b2 = b.building2
        cx2, cy2 = _cx_cy(getattr(b2, "results", None))
        if cx2: rows.append(["Upper", "X", f"{cx2:.4f}"])
        if cy2: rows.append(["Upper", "Y", f"{cy2:.4f}"])
    else:
        cx, cy = _cx_cy(getattr(b, "results", None))
        if cx: rows.append(["System", "X", f"{cx:.4f}"])
        if cy: rows.append(["System", "Y", f"{cy:.4f}"])

    if rows:
        _add_data_table(doc, headers, rows)
    else:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))


def _section_load_combinations(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("LOAD_COMBINATIONS", lang), level=1)
    if data.load_combinations is None or data.load_combinations.empty:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    df = data.load_combinations
    combo_names = df["Name"].unique()
    doc.add_paragraph(
        f"Total of {len(combo_names)} load combinations defined in the model."
    )

    # Build one-line format: "ComboName: SF1(Load1) + SF2(Load2) + ..."
    rows = []
    for name in combo_names:
        sub = df[df["Name"] == name]
        parts = []
        for _, r in sub.iterrows():
            sf = r.get("SF", 1)
            load = r.get("LoadName", r.get("Load", ""))
            try:
                sf_val = float(sf)
                if sf_val == 1.0:
                    parts.append(str(load))
                elif sf_val == -1.0:
                    parts.append(f"-{load}")
                else:
                    parts.append(f"{sf_val:g}({load})")
            except (ValueError, TypeError):
                parts.append(f"{sf}({load})")
        formula = " + ".join(parts).replace("+ -", "- ")
        rows.append([name, formula])

    table = _add_data_table(doc, ["Combination", "Definition"], rows)
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.65)
    _format_table_font(table, body_size=8.5, header_size=9.0)


def _section_story_forces(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("STORY_FORCES", lang), level=1)
    if not data.story_forces_data or not data.story_forces_fields:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Story forces from seismic load patterns with percentage of "
        "total base shear at each story."
    )
    _add_data_table(doc, data.story_forces_fields, data.story_forces_data)


def _section_drift(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("STORY_DRIFT", lang), level=1)
    if data.drift_data is None or data.drift_data.empty:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Maximum story drift values compared against code-allowable limits. "
        "Drifts exceeding the limit are flagged."
    )
    _add_df_table(doc, data.drift_data)

    if "Max Drift" in data.drift_data.columns and "Allowable Drift" in data.drift_data.columns:
        try:
            max_d = data.drift_data["Max Drift"].astype(float)
            allow = data.drift_data["Allowable Drift"].astype(float)
            if (max_d <= allow).all():
                doc.add_paragraph("Result: ALL DRIFT CHECKS PASSED.")
            else:
                doc.add_paragraph("Result: SOME DRIFT CHECKS FAILED — review required.")
        except Exception:
            pass


def _section_torsion(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("TORSIONAL_IRREGULARITY", lang), level=1)
    if data.torsion_data is None or data.torsion_data.empty:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Ratio of maximum diaphragm drift to average drift. "
        "A ratio > 1.2 indicates torsional irregularity; "
        "> 1.4 indicates extreme torsional irregularity."
    )
    _add_df_table(doc, data.torsion_data)

    if "Ratio" in data.torsion_data.columns:
        try:
            max_ratio = data.torsion_data["Ratio"].astype(float).max()
            if max_ratio <= 1.2:
                doc.add_paragraph("Result: No torsional irregularity detected.")
            elif max_ratio <= 1.4:
                doc.add_paragraph(
                    f"Result: Torsional irregularity detected (max ratio = {max_ratio:.3f})."
                )
            else:
                doc.add_paragraph(
                    f"Result: EXTREME torsional irregularity (max ratio = {max_ratio:.3f})."
                )
        except Exception:
            pass


def _section_pmm_columns(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("PMM_COLUMNS", lang), level=1)
    if data.pmm_data is None or data.pmm_data.empty:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Concrete column design results — PMM interaction ratio. "
        "Columns with ratio > 1.0 are overstressed."
    )
    _add_df_table(doc, data.pmm_data)

    if "PMMRatio" in data.pmm_data.columns:
        try:
            max_pmm = data.pmm_data["PMMRatio"].astype(float).max()
            if max_pmm <= 1.0:
                doc.add_paragraph(
                    f"Result: All columns adequate (max PMM ratio = {max_pmm:.3f})."
                )
            else:
                doc.add_paragraph(
                    f"Result: OVERSTRESSED columns found (max PMM ratio = {max_pmm:.3f})."
                )
        except Exception:
            pass


def _section_story_plans(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("STORY_PLANS", lang), level=1)
    if not data.story_plan_images:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Beam and column plan for each story with section names annotated."
    )
    for story in data.stories:
        img = data.story_plan_images.get(story)
        if img:
            doc.add_heading(f"Story: {story}", level=2)
            # Use full page width so all elements are legible
            _add_image(doc, img, width_inches=6.3)


def _section_area_loads(doc: Document, data: ReportData, lang: str):
    doc.add_heading(get_string("AREA_LOAD_PLANS", lang), level=1)
    if not data.area_load_images:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Floor area load plans for each story. Areas with identical "
        "load patterns are grouped and colour-coded."
    )
    for story in data.stories:
        img = data.area_load_images.get(story)
        if img:
            doc.add_heading(f"Story: {story}", level=2)
            _add_image(doc, img, width_inches=6.3)

            # Per-story load set table — sorted by pattern count (ascending)
            # so simpler load sets appear first and comparison is easy.
            story_areas = data.area_data.get(story, [])
            story_sets = sorted({a.load_set for a in story_areas})
            if story_sets and data.load_set_defs:
                headers = ["Load Set", "Load Patterns (kg/m\u00b2)"]
                # Priority order for load pattern display inside each set
                _PATTERN_ORDER = [
                    "DEAD", "DL", "SDL", "SDEAD", "PARTITION",
                    "LIVE", "LL", "LRED", "PARKING",
                    "LROOF", "ROOF",
                    "MASS",
                    "SNOW", "S",
                    "EV", "SEISMIC",
                ]

                def _pattern_sort_key(p: str) -> tuple:
                    pl = p.upper()
                    for rank, known in enumerate(_PATTERN_ORDER):
                        if pl.startswith(known):
                            return (rank, pl)
                    return (len(_PATTERN_ORDER), pl)

                def _fmt_loads(lsd) -> str:
                    if not lsd:
                        return ""
                    items = sorted(lsd.loads.items(), key=lambda kv: _pattern_sort_key(kv[0]))
                    return ", ".join(f"{p} = {v:.0f}" for p, v in items)

                # Sort load sets: fewer load patterns first
                def _set_sort_key(sname: str):
                    lsd = data.load_set_defs.get(sname)
                    n = len(lsd.loads) if lsd else 0
                    return (n, sname)

                rows = []
                for sname in sorted(story_sets, key=_set_sort_key):
                    lsd = data.load_set_defs.get(sname)
                    rows.append([sname, _fmt_loads(lsd)])
                _add_data_table(doc, headers, rows)


def _section_joint_shear(doc: Document, data: ReportData, lang: str):
    doc.add_heading("Joint Shear Check", level=1)
    if data.joint_shear_data is None or data.joint_shear_data.empty:
        doc.add_paragraph(get_string("NOT_AVAILABLE", lang))
        return

    doc.add_paragraph(
        "Beam-column joint shear check results."
    )
    _add_df_table(doc, data.joint_shear_data)


# ═══════════════════════════════════════════════════════════════════════════
# Section dispatch
# ═══════════════════════════════════════════════════════════════════════════

_SECTION_DISPATCH = {
    "model_settings":         _section_model_settings,
    "project_info":           _section_project_info,
    "structural_system":      _section_structural_system,
    "seismic_params":         _section_seismic_params,
    "earthquake_formulation": _section_earthquake_formulation,
    "earthquake_values":      _section_earthquake_values,
    "load_combinations":      _section_load_combinations,
    "story_forces":           _section_story_forces,
    "drift":                  _section_drift,
    "torsion":                _section_torsion,
    "joint_shear":            _section_joint_shear,
    "pmm_columns":            _section_pmm_columns,
    "story_plans":            _section_story_plans,
    "area_loads":             _section_area_loads,
}


# ═══════════════════════════════════════════════════════════════════════════
# High-level builder
# ═══════════════════════════════════════════════════════════════════════════

def create_docx_report(
    data: ReportData,
    config: ReportConfig | None = None,
    output_path: str | Path | None = None,
) -> Path:
    """Generate a full DOCX report from extracted data.

    Parameters
    ----------
    data : ReportData
        All report data (from ``extract_report_data``).
    config : ReportConfig, optional
        Report configuration.
    output_path : str | Path, optional
        Output file path.
    """
    config = config or ReportConfig(language="en")
    lang = config.language if config.language != "both" else "en"

    if output_path is None:
        stem = (data.project_name or "report").replace(" ", "_")
        output_path = Path(f"{stem}_report.docx")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists():
        try:
            with output_path.open("ab"):
                pass
        except PermissionError as exc:
            raise PermissionError(
                f"Cannot write the Word report because the file is open or locked: {output_path}\n"
                f"Please close it in Word and try again."
            ) from exc

    doc = _create_styled_doc(config)

    # ── Title page ────────────────────────────────────────────────────
    title = data.project_name or "Structural Report"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(30, 60, 120)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(get_string("REPORT_TITLE", lang))
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(get_string("GENERATED_BY", lang))
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(140, 140, 140)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────
    if config.include_table_of_contents:
        doc.add_heading(get_string("TABLE_OF_CONTENTS", lang), level=1)
        _add_table_of_contents(doc)
        doc.add_page_break()

    # ── Filter load combinations to design-active only ────────────────
    if (getattr(config, "filter_active_combinations", True)
            and data.design_combo_names
            and data.load_combinations is not None
            and not data.load_combinations.empty):
        mask = data.load_combinations["Name"].isin(data.design_combo_names)
        if mask.any():
            data = data.__class__(**{
                **{f.name: getattr(data, f.name) for f in data.__dataclass_fields__.values()},
                "load_combinations": data.load_combinations[mask].reset_index(drop=True),
            })

    # ── Sections ──────────────────────────────────────────────────────
    for section_key in config.active_sections:
        renderer = _SECTION_DISPATCH.get(section_key)
        if renderer:
            renderer(doc, data, lang)
            doc.add_page_break()
        else:
            doc.add_heading(config.get_section_name(section_key), level=1)
            doc.add_paragraph(f"[Section: {section_key} — content pending]")
            doc.add_page_break()

    # ── Save ──────────────────────────────────────────────────────────
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════

def _building_to_seismic_dict(building, direction: str) -> dict | None:
    """Convert Building attributes to a seismic_params dict
    compatible with ``latex_str.full_earthquake_calculation``.
    """
    if building is None:
        return None

    d = direction.lower()
    sys_attr = f"{d}_system"
    sys = getattr(building, sys_attr, None)

    acc = getattr(building, "acc", 0)

    sp = {
        "soil_type": getattr(building, "soil_type", ""),
        "A": acc,
        "I": getattr(building, "importance_factor", 1),
        "zone": getattr(building, "risk_level", ""),
        "risk_level": 3 if acc >= 0.3 else (2 if acc >= 0.25 else 1),
        "T0": getattr(getattr(building, "soil_properties", None), "T0", 0),
        "Ts": getattr(getattr(building, "soil_properties", None), "Ts", 0),
        "S": getattr(getattr(building, "soil_properties", None), "S", 0),
        "S0": getattr(getattr(building, "soil_properties", None), "S0", 0),
        "H": getattr(building, "height", 0),
        "height": getattr(building, "height", 0),
        "is_infill": getattr(building, "is_infill", False),
    }

    # System parameters
    if sys:
        sp["alpha"] = getattr(sys, "alpha", 0)
        sp["Ct"] = getattr(sys, "alpha", 0)
        sp["beta"] = getattr(sys, "pow", 0.75)
        sp[f"R{d}"] = getattr(sys, "Ru", 0)

    # Period values
    sp[f"T{d}"] = getattr(building, f"t{d}", 0)           # design period
    sp[f"T{d}_exp"] = getattr(building, f"t{d}_exp", 0)    # empirical
    sp[f"T{d}_an"] = getattr(building, f"t{d}_an", 0)      # analytical
    sp[f"T{d}_design"] = getattr(building, f"t{d}", 0)     # design = t{d}

    # Reflection factor components from the ReflectionFactor object
    refl = getattr(building, f"soil_reflection_prop_{d}", None)
    if refl:
        sp[f"B1{d}"] = getattr(refl, "B1", 0)
        sp[f"N{d}"] = getattr(refl, "N", 1)
    else:
        sp[f"B1{d}"] = 0
        sp[f"N{d}"] = 1

    sp[f"B{d}"] = getattr(building, f"b{d}", 0)
    sp[f"K{d}"] = getattr(building, f"k{d}", 0)
    sp[f"C{d}"] = 0

    if hasattr(building, "results") and building.results and building.results[0] is True:
        _, cx, cy = building.results
        sp["Cx"] = cx
        sp["Cy"] = cy
        sp[f"C{d}"] = cx if d == "x" else cy

    return sp


def _building_all_to_seismic_dict(building, direction: str) -> dict | None:
    """Build seismic params dict using the *combined* period for dual systems.

    Uses ``tx_all`` / ``ty_all`` (combined period from both portions) and the
    lower-system structural parameters so the report can show the governing
    base-shear calculation for the overall building.

    Returns None if ``building.building2`` does not exist (single system).
    """
    if building is None or not hasattr(building, "building2") or building.building2 is None:
        return None

    d = direction.lower()
    sys = getattr(building, f"{d}_system", None)
    acc = getattr(building, "acc", 0)

    sp = {
        "soil_type": getattr(building, "soil_type", ""),
        "A": acc,
        "I": getattr(building, "importance_factor", 1),
        "zone": getattr(building, "risk_level", ""),
        "risk_level": 3 if acc >= 0.3 else (2 if acc >= 0.25 else 1),
        "T0": getattr(getattr(building, "soil_properties", None), "T0", 0),
        "Ts": getattr(getattr(building, "soil_properties", None), "Ts", 0),
        "S": getattr(getattr(building, "soil_properties", None), "S", 0),
        "S0": getattr(getattr(building, "soil_properties", None), "S0", 0),
        # Use full combined height (lower + upper)
        "H": getattr(building, "height", 0) + getattr(building.building2, "height", 0),
        "height": getattr(building, "height", 0) + getattr(building.building2, "height", 0),
        "is_infill": getattr(building, "is_infill", False),
    }

    if sys:
        sp["alpha"] = getattr(sys, "alpha", 0)
        sp["Ct"] = getattr(sys, "alpha", 0)
        sp["beta"] = getattr(sys, "pow", 0.75)
        sp[f"R{d}"] = getattr(sys, "Ru", 0)

    # Use combined-period values
    t_all = getattr(building, f"t{d}_all", getattr(building, f"t{d}", 0))
    sp[f"T{d}"] = t_all
    sp[f"T{d}_exp"] = getattr(building, f"t{d}_exp_all", getattr(building, f"t{d}_exp", 0))
    sp[f"T{d}_an"] = getattr(building, f"t{d}_an_all", getattr(building, f"t{d}_an", 0))
    sp[f"T{d}_design"] = t_all

    refl_all = getattr(building, f"soil_reflection_prop_all_{d}", None)
    if refl_all:
        sp[f"B1{d}"] = getattr(refl_all, "B1", 0)
        sp[f"N{d}"] = getattr(refl_all, "N", 1)
    else:
        sp[f"B1{d}"] = 0
        sp[f"N{d}"] = 1

    sp[f"B{d}"] = getattr(building, f"b{d}_all", 0)
    sp[f"K{d}"] = getattr(building, f"k{d}_all", 0)
    sp[f"C{d}"] = 0

    res = getattr(building, f"results_all_bot", None)
    if res and res[0] is True:
        _, cx, cy = res
        sp["Cx"] = cx
        sp["Cy"] = cy
        sp[f"C{d}"] = cx if d == "x" else cy

    return sp

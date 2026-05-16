# pyright: reportMissingImports=false, reportMissingModuleSource=false

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterable

import pandas as pd

__all__ = ["ControlsInputReportExporter"]


class ControlsInputReportExporter:
    """Export controls-input results to Excel, Word, and PDF."""

    def __init__(self, results: dict[str, dict[str, Any]]):
        self.results = results

    def summary_dataframe(self) -> pd.DataFrame:
        rows = []
        for key, result in self.results.items():
            summary = result.get("summary", {})
            rows.append(
                {
                    "Control Key": key,
                    "Control Name": result.get("control_name", key),
                    "Status": result.get("status", ""),
                    "Checked": summary.get("total_checked", 0),
                    "Failed": summary.get("failed", 0),
                    "Passed": summary.get("passed", 0),
                }
            )
        return pd.DataFrame(rows)

    def details_dataframe(self) -> pd.DataFrame:
        rows: list[dict[str, Any]] = []         
        for key, result in self.results.items():
            for detail in result.get("details", []):
                row: dict[str, Any] = {
                    "Control Key": key,
                    "Control Name": result.get("control_name", key),
                    "Status": result.get("status", ""),
                }
                for field, value in detail.items():
                    # Skip message and expected fields
                    if field in ("message", "expected"):
                        continue
                    if field == "actual":
                        if isinstance(value, dict):
                            # Expand dict keys directly, e.g. i22, i33, mass
                            for sub_key, sub_val in value.items():
                                row[sub_key] = sub_val
                        else:
                            row["actual"] = value
                    elif isinstance(value, dict):
                        for sub_key, sub_val in value.items():
                            row[f"{field}_{sub_key}"] = sub_val
                    else:
                        row[field] = value
                rows.append(row)
        return pd.DataFrame(rows)

    def export_excel(self, file_path: str | Path) -> Path:
        target = Path(file_path)
        with pd.ExcelWriter(target) as writer:
            self.summary_dataframe().to_excel(writer, sheet_name="Summary", index=False)
            self.details_dataframe().to_excel(writer, sheet_name="Details", index=False)
        return target

    def export_word(self, file_path: str | Path) -> Path:
        from docx import Document

        target = Path(file_path)
        document = Document()
        document.add_heading("Input Controls Report", level=1)

        summary_df = self.summary_dataframe()
        document.add_heading("Summary", level=2)
        self._append_table(document, summary_df)

        details_df = self.details_dataframe()
        if not details_df.empty:
            document.add_heading("Details", level=2)
            self._append_table(document, details_df)

        document.save(target)
        return target

    def export_pdf(self, file_path: str | Path) -> Path:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        target = Path(file_path)
        doc = SimpleDocTemplate(str(target), pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        story = [Paragraph("Input Controls Report", styles["Title"]), Spacer(1, 12)]

        summary_df = self.summary_dataframe()
        if not summary_df.empty:
            story.append(Paragraph("Summary", styles["Heading2"]))
            story.append(self._build_pdf_table(summary_df, colors, Table, TableStyle))
            story.append(Spacer(1, 12))

        details_df = self.details_dataframe()
        if not details_df.empty:
            story.append(Paragraph("Details", styles["Heading2"]))
            story.append(self._build_pdf_table(details_df.head(200), colors, Table, TableStyle))

        doc.build(story)
        return target

    @staticmethod
    def _append_table(document: Any, df: pd.DataFrame) -> None:
        if df.empty:
            document.add_paragraph("No data available.")
            return
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        for index, column in enumerate(df.columns):
            table.rows[0].cells[index].text = str(column)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for index, column in enumerate(df.columns):
                cells[index].text = str(row[column])

    @staticmethod
    def _build_pdf_table(df: pd.DataFrame, colors: Any, table_cls: Any, style_cls: Any) -> Any:
        data = [list(df.columns)] + df.astype(str).values.tolist()
        table = table_cls(data, repeatRows=1)
        table.setStyle(
            style_cls(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244061")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
                ]
            )
        )
        return table
